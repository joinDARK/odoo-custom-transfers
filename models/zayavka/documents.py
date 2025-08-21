from odoo import models
import base64
import io
import logging
import re
import zipfile
import os
import tempfile
import subprocess
from xml.etree import ElementTree as ET

try:
    import pymupdf
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False



_logger = logging.getLogger(__name__)


class AmanatZayavkaDocuments(models.Model):
    _inherit = 'amanat.zayavka'

    def extract_text_from_docx_attachment(self, attachment):
        try:
            if not attachment or not attachment.datas:
                _logger.warning("[extract_text_from_docx_attachment] пустое вложение или отсутствуют данные")
                return None

            # Быстрая проверка расширения/типов
            name_lower = (attachment.name or '').lower()
            mimetype = (attachment.mimetype or '').lower()
            if not (name_lower.endswith('.docx') or 'officedocument.wordprocessingml.document' in mimetype):
                _logger.info(f"[extract_text_from_docx_attachment] вложение {attachment.name} не DOCX (mimetype={mimetype})")
                return None

            docx_bytes = base64.b64decode(attachment.datas)
            return self._extract_text_from_docx_bytes(docx_bytes)

        except Exception as e:
            _logger.error(f"[extract_text_from_docx_attachment] ошибка обработки {attachment and attachment.name}: {str(e)}")
            return None

    @staticmethod
    def _extract_text_from_docx_bytes(docx_bytes):
        if not docx_bytes:
            return None

        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                namelist = zf.namelist()

                # Основные и дополнительные части, из которых можем извлечь текст
                primary_parts = [
                    'word/document.xml',
                ]
                # Заголовки/подвалы/сноски, если есть
                extra_parts = [n for n in namelist if n.startswith('word/header') and n.endswith('.xml')]
                extra_parts += [n for n in namelist if n.startswith('word/footer') and n.endswith('.xml')]
                for part in ['word/footnotes.xml', 'word/endnotes.xml']:
                    if part in namelist:
                        extra_parts.append(part)

                parts_to_read = [p for p in primary_parts if p in namelist] + extra_parts
                if not parts_to_read:
                    # Нестандартный DOCX без ожидаемых частей
                    return None

                all_paragraphs = []
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                def extract_paragraphs(xml_content):
                    try:
                        root = ET.fromstring(xml_content)
                    except ET.ParseError:
                        return []
                    paragraphs = []
                    # Параграфы <w:p>
                    for p in root.findall('.//w:p', ns):
                        text_chunks = []
                        # Пробегаемся по run-ам <w:r>
                        for node in p.iter():
                            tag = node.tag.split('}')[-1]
                            if tag == 't':  # <w:t>
                                if node.text:
                                    text_chunks.append(node.text)
                            elif tag == 'br':  # <w:br/>
                                text_chunks.append('\n')
                            elif tag == 'tab':  # <w:tab/>
                                text_chunks.append('\t')
                        paragraph_text = ''.join(text_chunks)
                        paragraph_text = paragraph_text.replace('\r', '')
                        # Нормализуем множественные пробелы внутри строки
                        paragraph_text = re.sub(r'[ \t\xa0]+', ' ', paragraph_text)
                        paragraph_text = paragraph_text.strip()
                        if paragraph_text:
                            paragraphs.append(paragraph_text)
                    return paragraphs

                for part_name in parts_to_read:
                    try:
                        with zf.open(part_name) as fp:
                            xml_data = fp.read()
                        all_paragraphs.extend(extract_paragraphs(xml_data))
                    except KeyError:
                        # Части может не быть — пропускаем
                        continue

                if not all_paragraphs:
                    return None

                # Склеиваем с сохранением абзацев
                text = '\n'.join(all_paragraphs)
                # Чистим лишние пустые строки
                text = re.sub(r'\n{3,}', '\n\n', text)
                
                # Логируем размер извлеченного текста для диагностики
                if text and text.strip():
                    final_text = text.strip()
                    _logger.info(f"[_extract_text_from_docx_bytes] Извлечен текст: {len(final_text)} символов, {len(all_paragraphs)} параграфов")
                    _logger.info(f"[_extract_text_from_docx_bytes] Первые 200 символов: {final_text[:200]}...")
                    if len(final_text) > 400:
                        _logger.info(f"[_extract_text_from_docx_bytes] Последние 200 символов: ...{final_text[-200:]}")
                    return final_text
                else:
                    _logger.warning("[_extract_text_from_docx_bytes] Извлеченный текст пуст")
                    return None

        except zipfile.BadZipFile:
            # Файл не является корректным DOCX
            return None
        except Exception:
            return None

    # === New helpers for auto-signing Assignment Start with STELLAR ===
    def _detect_file_type(self, file_data, file_name=None):
        """Lightweight file type detection for pdf/docx/zip."""
        if not file_data:
            _logger.warning(f"[_detect_file_type] No file data provided")
            return None
        try:
            if isinstance(file_data, str):
                decoded = base64.b64decode(file_data)
            else:
                decoded = file_data

            _logger.info(f"[_detect_file_type] File name: {file_name}, first 10 bytes: {decoded[:10]}")

            if decoded[:4] == b'%PDF':
                _logger.info(f"[_detect_file_type] Detected as PDF by header")
                return 'pdf'
            
            # Дополнительная проверка по расширению файла для PDF
            if file_name and file_name.lower().endswith('.pdf'):
                _logger.info(f"[_detect_file_type] Detected as PDF by extension (fallback)")
                return 'pdf'
            if decoded[:2] == b'PK':
                _logger.info(f"[_detect_file_type] Detected ZIP-based file")
                # Rely on extension hint if provided
                if file_name and file_name.lower().endswith('.docx'):
                    _logger.info(f"[_detect_file_type] Detected as DOCX by extension")
                    return 'docx'
                # Inspect minimal ZIP content header
                if b'word/' in decoded[:4000]:
                    _logger.info(f"[_detect_file_type] Detected as DOCX by content")
                    return 'docx'
                _logger.info(f"[_detect_file_type] Detected as generic ZIP")
                return 'zip'
            
            _logger.warning(f"[_detect_file_type] Unknown file type, first 10 bytes: {decoded[:10]}")
            return 'unknown'
        except Exception as e:
            _logger.error(f"[_detect_file_type] Error detecting file type: {e}")
            return 'unknown'

    def _fix_table_formatting(self, docx_path):
        """Исправляет форматирование таблиц в DOCX для лучшей конвертации в PDF"""
        if not PYTHON_DOCX_AVAILABLE:
            _logger.warning("python-docx недоступен, пропускаем исправление форматирования таблиц")
            return docx_path
            
        try:
            _logger.info("🔧 Исправляем форматирование таблиц...")
            doc = Document(docx_path)
            
            # Обрабатываем все таблицы в документе
            for table in doc.tables:
                # Устанавливаем фиксированную ширину таблицы
                table.autofit = False
                
                for row in table.rows:
                    # Устанавливаем минимальную высоту строки
                    row.height = Inches(0.3)
                    
                    for cell in row.cells:
                        # Выравнивание по верхнему краю для всех ячеек
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                        
                        # Устанавливаем отступы в ячейках
                        cell.margin_top = Inches(0.05)
                        cell.margin_bottom = Inches(0.05)
                        cell.margin_left = Inches(0.1)
                        cell.margin_right = Inches(0.1)
                        
                        # Обрабатываем параграфы в ячейке
                        for paragraph in cell.paragraphs:
                            # Устанавливаем интервал между строками
                            paragraph.paragraph_format.line_spacing = 1.0
                            paragraph.paragraph_format.space_before = 0
                            paragraph.paragraph_format.space_after = 0
            
            # Сохраняем исправленный документ
            fixed_path = docx_path.replace('.docx', '_fixed.docx')
            doc.save(fixed_path)
            _logger.info("✅ Форматирование таблиц исправлено")
            return fixed_path
            
        except Exception as e:
            _logger.warning(f"⚠️ Не удалось исправить форматирование таблиц: {e}")
            return docx_path

    def _convert_docx_to_pdf(self, docx_data):
        """Convert DOCX (base64 or bytes) to PDF using LibreOffice with preserved formatting."""
        if not docx_data:
            raise ValueError('DOCX data is empty')

        try:
            raw = base64.b64decode(docx_data) if isinstance(docx_data, str) else docx_data
            
            # Создаем временный файл с правильным именем (важно для LibreOffice)
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, prefix='odoo_convert_') as docx_temp:
                docx_temp.write(raw)
                docx_path = docx_temp.name
            
            # Исправляем форматирование таблиц перед конвертацией
            fixed_docx_path = self._fix_table_formatting(docx_path)

            out_dir = tempfile.mkdtemp()
            try:
                # Оптимизированная команда LibreOffice для сохранения форматирования
                _logger.info(f"Конвертируем DOCX в PDF: {fixed_docx_path}")
                
                # Метод 1: Точная конвертация с сохранением форматирования таблиц
                cmd_precise = [
                    'libreoffice', '--headless', '--invisible', '--nodefault', '--nolockcheck',
                    '--convert-to', 'pdf:writer_pdf_Export:{"EmbedStandardFonts":true,"ExportFormFields":false,"FormsType":0,"ExportBookmarks":false,"ExportNotes":false,"ExportNotesPages":false,"ExportOnlyNotesPages":false,"UseTransitionEffects":false,"IsSkipEmptyPages":false,"IsAddStream":false,"SelectPdfVersion":0,"CompressMode":2,"JPEGQuality":90,"ImageResolution":300}',
                    '--outdir', out_dir,
                    fixed_docx_path
                ]
                
                # Метод 2: Простая конвертация (fallback)
                cmd_simple = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', out_dir, fixed_docx_path
                ]
                
                # Пробуем точную конвертацию
                result = None
                try:
                    _logger.info("Пробуем точную конвертацию...")
                    result = subprocess.run(cmd_precise, capture_output=True, text=True, timeout=90)
                    if result.returncode == 0:
                        _logger.info("✅ Точная конвертация успешна!")
                    else:
                        _logger.warning(f"Точная конвертация не удалась: {result.stderr}")
                        raise RuntimeError("Fallback to simple conversion")
                except Exception as e:
                    _logger.warning(f"Ошибка точной конвертации: {e}, переходим на простую...")
                    result = subprocess.run(cmd_simple, capture_output=True, text=True, timeout=60)
                    if result.returncode == 0:
                        _logger.info("✅ Простая конвертация успешна!")
                    else:
                        _logger.error(f"❌ Простая конвертация тоже не удалась: {result.stderr}")
                if result.returncode != 0:
                    raise RuntimeError(f'LibreOffice error: {result.stderr}')

                pdf_filename = os.path.splitext(os.path.basename(fixed_docx_path))[0] + '.pdf'
                pdf_path = os.path.join(out_dir, pdf_filename)
                if not os.path.exists(pdf_path):
                    raise FileNotFoundError('Converted PDF not found')

                with open(pdf_path, 'rb') as f:
                    pdf_bytes = f.read()
                return base64.b64encode(pdf_bytes)
            finally:
                # Очищаем временные файлы
                try:
                    os.unlink(docx_path)
                except Exception:
                    pass
                try:
                    if fixed_docx_path != docx_path:
                        os.unlink(fixed_docx_path)
                except Exception:
                    pass
                try:
                    pdf_filename = os.path.splitext(os.path.basename(fixed_docx_path))[0] + '.pdf'
                    pdf_path_cleanup = os.path.join(out_dir, pdf_filename)
                    if os.path.exists(pdf_path_cleanup):
                        os.unlink(pdf_path_cleanup)
                    os.rmdir(out_dir)
                except Exception:
                    pass
        except Exception as e:
            _logger.error(f"DOCX→PDF conversion failed: {e}")
            raise



    def _find_text_positions_in_pdf(self, pdf_bytes, keywords):
        """
        Find positions of any of the keywords (case-insensitive) in the PDF.
        Returns list of dicts: {page_number (1-based), bbox: (x0,y0,x1,y1), text}
        """
        if not PYMUPDF_AVAILABLE:
            return []
        if not pdf_bytes:
            return []
        matches = []
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            lowered = [kw.lower() for kw in keywords]
            for page_index in range(len(doc)):
                page = doc[page_index]
                text_dict = page.get_text('dict')
                for block in text_dict.get('blocks', []):
                    for line in block.get('lines', []):
                        for span in line.get('spans', []):
                            span_text = (span.get('text') or '').strip()
                            if not span_text:
                                continue
                            if any(kw in span_text.lower() for kw in lowered):
                                bbox = span.get('bbox')
                                if bbox:
                                    matches.append({
                                        'page_number': page_index + 1,
                                        'bbox': tuple(bbox),
                                        'text': span_text,
                                    })
            doc.close()
        except Exception as e:
            _logger.error(f"Error searching PDF text: {e}")
        return matches

    def _find_signature_position_in_pdf(self, pdf_bytes):
        """
        Ищет правильное место для подписи в PDF документе.
        Ищет строку "Субагент/ Subagent" в конце документа (обычно в области подписей).
        Возвращает список позиций для подписи.
        """
        if not PYMUPDF_AVAILABLE or not pdf_bytes:
            return []
        
        matches = []
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            
            # Ищем паттерны, которые указывают на место подписи
            signature_patterns = [
                'Субагент/ Subagent',
                'Субагент/Subagent', 
                'Субагент / Subagent',
                'Субагент Subagent',
                'Субагент:', 
                'Subagent:'
            ]
            
            for page_index in range(len(doc)):
                page = doc[page_index]
                text_dict = page.get_text('dict')
                
                for block in text_dict.get('blocks', []):
                    for line in block.get('lines', []):
                        # Собираем весь текст строки
                        line_text = ''
                        line_bbox = None
                        
                        for span in line.get('spans', []):
                            span_text = (span.get('text') or '').strip()
                            if span_text:
                                line_text += span_text + ' '
                                if line_bbox is None:
                                    line_bbox = span.get('bbox')
                        
                        line_text = line_text.strip()
                        
                        # Проверяем, содержит ли строка один из паттернов подписи
                        for pattern in signature_patterns:
                            if pattern.lower() in line_text.lower():
                                _logger.info(f"[_find_signature_position_in_pdf] Found signature pattern '{pattern}' in line: '{line_text}'")
                                if line_bbox:
                                    matches.append({
                                        'page_number': page_index + 1,
                                        'bbox': tuple(line_bbox),
                                        'text': line_text,
                                        'pattern': pattern
                                    })
                                break
            
            doc.close()
            _logger.info(f"[_find_signature_position_in_pdf] Found {len(matches)} signature positions")
            
        except Exception as e:
            _logger.error(f"[_find_signature_position_in_pdf] Error searching for signature position: {e}")
        
        return matches

    def _detect_signature_type_from_pdf(self, pdf_bytes):
        """
        Определяет тип подписи (ТДК или СТЕЛЛАР) по тексту в PDF документе.
        Ищет упоминания 'ТДК', 'TDK', 'СТЕЛЛАР', 'STELLAR' рядом с местом подписи.
        """
        if not PYMUPDF_AVAILABLE or not pdf_bytes:
            _logger.warning("[_detect_signature_type_from_pdf] PyMuPDF not available or no PDF data")
            return 'СТЕЛЛАР'
        
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            
            tdk_keywords = ['ТДК', 'TDK', 'тдк', 'Тдк']
            stellar_keywords = ['СТЕЛЛАР', 'STELLAR', 'стеллар', 'Стеллар']
            
            all_text = ""
            for page_index in range(len(doc)):
                page = doc[page_index]
                page_text = page.get_text()
                all_text += page_text.lower()
            
            doc.close()
            
            tdk_found = any(keyword.lower() in all_text for keyword in tdk_keywords)
            stellar_found = any(keyword.lower() in all_text for keyword in stellar_keywords)
            
            _logger.info(f"[_detect_signature_type_from_pdf] TDK found: {tdk_found}, STELLAR found: {stellar_found}")
            
            if tdk_found and stellar_found:
                tdk_count = sum(all_text.count(keyword.lower()) for keyword in tdk_keywords)
                stellar_count = sum(all_text.count(keyword.lower()) for keyword in stellar_keywords)
                _logger.info(f"[_detect_signature_type_from_pdf] TDK count: {tdk_count}, STELLAR count: {stellar_count}")
                return 'ТДК' if tdk_count > stellar_count else 'СТЕЛЛАР'
            elif tdk_found:
                return 'ТДК'
            elif stellar_found:
                return 'СТЕЛЛАР'
            else:
                _logger.info("[_detect_signature_type_from_pdf] No signature type keywords found, defaulting to СТЕЛЛАР")
                return 'СТЕЛЛАР'
                
        except Exception as e:
            _logger.error(f"[_detect_signature_type_from_pdf] Error detecting signature type: {e}")
            return 'СТЕЛЛАР'

    def _find_signature_and_stamp_records(self, signature_type):
        """
        Находит записи подписи и печати для указанного типа (ТДК или СТЕЛЛАР).
        Возвращает кортеж (signature_record, stamp_record).
        """
        if signature_type == 'ТДК':
            search_terms = ['ТДК', 'TDK']
        else:  # СТЕЛЛАР
            search_terms = ['СТЕЛЛАР', 'STELLAR']
        
        sig_record = None
        stamp_record = None
        
        # Ищем подпись
        for term in search_terms:
            sig_record = self.env['signature.library'].search([
                ('name', 'ilike', term),
                ('signature_type', '=', 'signature'),
                ('active', '=', True)
            ], limit=1)
            if sig_record:
                break
        
        # Ищем печать
        for term in search_terms:
            stamp_record = self.env['signature.library'].search([
                ('name', 'ilike', term),
                ('signature_type', '=', 'stamp'),
                ('active', '=', True)
            ], limit=1)
            if stamp_record:
                break
        
        _logger.info(f"[_find_signature_and_stamp_records] Found signature: {sig_record.name if sig_record else 'None'}")
        _logger.info(f"[_find_signature_and_stamp_records] Found stamp: {stamp_record.name if stamp_record else 'None'}")
        
        return sig_record, stamp_record

    def _sign_single_document(self, pdf_bytes, target_match, sig_record, stamp_record, signature_type):
        """
        Подписывает один PDF документ с учетом типа подписи (ТДК или СТЕЛЛАР).
        Возвращает байты подписанного DOCX или None в случае ошибки.
        """
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            page_index = target_match['page_number'] - 1
            
            if page_index >= len(doc):
                _logger.warning(f"[_sign_single_document] Page index out of range: {page_index}")
                doc.close()
                return None
                
            page = doc[page_index]
            x0, y0, x1, y1 = target_match['bbox']

            # Настройки размещения в зависимости от типа подписи
            if signature_type == 'ТДК':
                # Для ТДК: подпись меньше и выше
                sig_w = (sig_record.default_width or 150) * 1.1  # увеличиваем на 10% (меньше чем у СТЕЛЛАР)
                sig_h = (sig_record.default_height or 50) * 1.1   # увеличиваем на 10%
                padding_y = -12  # поднимаем выше (больший отрицательный отступ)
                _logger.info(f"[_sign_single_document] Using ТДК settings: size multiplier 1.1, padding_y: {padding_y}")
            else:  # СТЕЛЛАР
                # Для СТЕЛЛАР: как было (больше и ниже)
                sig_w = (sig_record.default_width or 150) * 1.3  # увеличиваем на 30%
                sig_h = (sig_record.default_height or 50) * 1.3   # увеличиваем на 30%
                padding_y = -4  # стандартный отступ
                _logger.info(f"[_sign_single_document] Using СТЕЛЛАР settings: size multiplier 1.3, padding_y: {padding_y}")

            # Размещение подписи
            sig_x = x0
            sig_y = y1 + padding_y  # размещаем относительно текста

            sig_rect = pymupdf.Rect(sig_x, sig_y, sig_x + sig_w, sig_y + sig_h)
            page.insert_image(sig_rect, stream=base64.b64decode(sig_record.image))

            # Размещение печати (одинаково для всех типов)
            stamp_w = stamp_record.default_width or 100  # размер печати по умолчанию
            stamp_h = stamp_record.default_height or 100  # размер печати по умолчанию
            padding_x = 15  # отступ между подписью и печатью
            stamp_x = sig_x + sig_w + padding_x
            stamp_y = sig_y  # выравниваем по верхнему краю подписи

            stamp_rect = pymupdf.Rect(stamp_x, stamp_y, stamp_x + stamp_w, stamp_y + stamp_h)
            page.insert_image(stamp_rect, stream=base64.b64decode(stamp_record.image))

            # Сохраняем подписанный PDF
            signed_pdf_bytes = doc.tobytes()
            doc.close()
            
            _logger.info(f"[_sign_single_document] Successfully signed document with {signature_type} signature - returning PDF")
            return {'pdf': signed_pdf_bytes}
            
        except Exception as e:
            _logger.error(f"[_sign_single_document] Error while signing document: {e}")
            if 'doc' in locals():
                doc.close()
            return None

    def should_auto_sign_document(self, attachment):
        """
        Быстрая проверка: нужно ли автоматически подписывать документ.
        Возвращает True только если документ содержит:
        1. Текст 'Субагент'/'Subagent' 
        2. Информацию о 'ТДК'/'TDK' или 'СТЕЛЛАР'/'STELLAR'
        """
        if not attachment or not attachment.datas:
            return False
            
        try:
            # Определяем тип файла
            file_type = self._detect_file_type(attachment.datas, attachment.name or '')
            if file_type == 'docx':
                pdf_b64 = self._convert_docx_to_pdf(attachment.datas)
                pdf_bytes = base64.b64decode(pdf_b64)
            elif file_type == 'pdf':
                pdf_bytes = base64.b64decode(attachment.datas)
            else:
                _logger.info(f"[should_auto_sign_document] Unsupported file type '{file_type}' for '{attachment.name}'")
                return False

            # УСЛОВИЕ 1: Проверяем наличие места для подписи 'Субагент/ Subagent'
            subagent_matches = self._find_signature_position_in_pdf(pdf_bytes)
            if not subagent_matches:
                _logger.info(f"[should_auto_sign_document] Document '{attachment.name}' does not contain signature position 'Субагент/ Subagent' - no signing needed")
                return False

            # УСЛОВИЕ 2: Проверяем наличие информации о ТДК или СТЕЛЛАР
            signature_type = self._detect_signature_type_from_pdf(pdf_bytes)
            if signature_type not in ['ТДК', 'СТЕЛЛАР']:
                _logger.info(f"[should_auto_sign_document] Document '{attachment.name}' does not contain ТДК/СТЕЛЛАР information - no signing needed")
                return False
            
            _logger.info(f"[should_auto_sign_document] Document '{attachment.name}' meets all conditions for auto-signing: Субагент + {signature_type}")
            return True
            
        except Exception as e:
            _logger.error(f"[should_auto_sign_document] Error checking document '{attachment.name}': {e}")
            return False

    def auto_sign_assignment_with_stellar(self):
        """
        When files are uploaded to `assignment_attachments`, automatically:
        - First check if document contains 'Субагент'/'Subagent' text
        - ONLY if text found: detect signature type (ТДК or СТЕЛЛАР) from document text
        - supports both PDF and DOCX input files
        - place appropriate signature slightly below that text
        - place appropriate stamp to the right of the signature
        - save signed PDF document
        Save signed PDF file to `assignment_end_attachments`.
        """
        _logger.info(f"[auto_sign_assignment_with_stellar] ===== STARTING AUTO-SIGNING PROCESS FOR RECORD {self.id} =====")
        self.ensure_one()
        attachments = self.assignment_attachments
        _logger.info(f"[auto_sign_assignment_with_stellar] Found {len(attachments)} attachments in assignment_attachments field")
        
        # Детальная информация о каждом вложении
        for i, att in enumerate(attachments):
            _logger.info(f"[auto_sign_assignment_with_stellar] Attachment {i+1}: ID={att.id}, name='{att.name}', size={len(att.datas) if att.datas else 0} bytes")
        
        if not attachments:
            _logger.info("[auto_sign_assignment_with_stellar] No attachments to process")
            return False

        if not PYMUPDF_AVAILABLE:
            _logger.warning("[auto_sign_assignment_with_stellar] PyMuPDF not available; skipping auto signing")
            return False

        # Очищаем существующие выходные вложения (выход 2)
        if self.assignment_end_attachments:
            self.assignment_end_attachments.unlink()

        signed_attachments = []
        success_count = 0

        # Обрабатываем каждый файл
        for idx, attachment in enumerate(attachments):
            _logger.info(f"[auto_sign_assignment_with_stellar] Processing attachment {idx+1}/{len(attachments)}: {attachment.name}")
            
            if not attachment.datas:
                _logger.warning(f"[auto_sign_assignment_with_stellar] Attachment '{attachment.name}' has no data, skipping")
                continue
            
            try:
                # Определяем тип файла
                file_type = self._detect_file_type(attachment.datas, attachment.name or '')
                if file_type == 'docx':
                    pdf_b64 = self._convert_docx_to_pdf(attachment.datas)
                    pdf_bytes = base64.b64decode(pdf_b64)
                elif file_type == 'pdf':
                    pdf_bytes = base64.b64decode(attachment.datas)
                else:
                    _logger.warning(f"[auto_sign_assignment_with_stellar] Unsupported file type '{file_type}' for '{attachment.name}', skipping")
                    continue

                # ПЕРВАЯ ПРОВЕРКА: ищем место для подписи - должно быть "Субагент/ Subagent" в конце документа
                matches = self._find_signature_position_in_pdf(pdf_bytes)
                if not matches:
                    _logger.info(f"[auto_sign_assignment_with_stellar] Document '{attachment.name}' does not contain signature position 'Субагент/ Subagent' - SKIPPING (не требует подписи)")
                    continue
                
                _logger.info(f"[auto_sign_assignment_with_stellar] Document '{attachment.name}' contains signature position 'Субагент/ Subagent' - proceeding with signing")

                # Определяем тип подписи
                signature_type = self._detect_signature_type_from_pdf(pdf_bytes)
                _logger.info(f"[auto_sign_assignment_with_stellar] Detected signature type for '{attachment.name}': {signature_type}")

                # Ищем подпись и печать
                sig_record, stamp_record = self._find_signature_and_stamp_records(signature_type)
                if not sig_record or not stamp_record:
                    _logger.warning(f"[auto_sign_assignment_with_stellar] {signature_type} signature/stamp not found for '{attachment.name}', skipping")
                    continue

                # Подписываем документ (возвращает только PDF)
                signed_document_data = self._sign_single_document(pdf_bytes, matches[0], sig_record, stamp_record, signature_type)
                if signed_document_data and signed_document_data.get('pdf'):
                    base_name = attachment.name.rsplit('.', 1)[0]
                    
                    # Создаем подписанный PDF файл
                    pdf_signed_name = f"signed_{base_name}_{idx+1}.pdf"
                    pdf_attachment = self.env['ir.attachment'].create({
                        'name': pdf_signed_name,
                        'datas': base64.b64encode(signed_document_data['pdf']),
                'res_model': self._name,
                'res_id': self.id,
                'mimetype': 'application/pdf',
            })
                    signed_attachments.append(pdf_attachment.id)
                    success_count += 1
                    _logger.info(f"[auto_sign_assignment_with_stellar] Successfully signed '{attachment.name}' -> PDF: '{pdf_signed_name}'")

            except Exception as e:
                _logger.error(f"[auto_sign_assignment_with_stellar] Error processing '{attachment.name}': {e}")
                continue

        # Привязываем все подписанные файлы
        if signed_attachments:
            _logger.info(f"[auto_sign_assignment_with_stellar] Attempting to save {len(signed_attachments)} signed files to assignment_end_attachments")
            _logger.info(f"[auto_sign_assignment_with_stellar] Signed attachment IDs: {signed_attachments}")
            
            # Очищаем старые файлы и добавляем новые
            self.assignment_end_attachments = [(6, 0, signed_attachments)]
            
            # Проверяем что файлы действительно сохранились
            saved_attachments = self.assignment_end_attachments
            _logger.info(f"[auto_sign_assignment_with_stellar] After saving - assignment_end_attachments count: {len(saved_attachments)}")
            for att in saved_attachments:
                _logger.info(f"[auto_sign_assignment_with_stellar] Saved attachment: ID={att.id}, name='{att.name}'")
            
            _logger.info(f"[auto_sign_assignment_with_stellar] Successfully processed {success_count}/{len(attachments)} files")
            
            # Проставляем дату подписания поручения при успешном подписании
            from datetime import date
            today = date.today()
            
            # Обновляем поля даты подписания
            vals_to_update = {}
            if not self.instruction_signed_date:
                vals_to_update['instruction_signed_date'] = today
                _logger.info(f"[auto_sign_assignment_with_stellar] Setting instruction_signed_date to {today}")
            
            if not self.assignment_signed_sovcom:
                vals_to_update['assignment_signed_sovcom'] = today
                _logger.info(f"[auto_sign_assignment_with_stellar] Setting assignment_signed_sovcom to {today}")
            
            if vals_to_update:
                self.write(vals_to_update)
                _logger.info(f"[auto_sign_assignment_with_stellar] Updated signature date fields: {list(vals_to_update.keys())}")
            else:
                _logger.info(f"[auto_sign_assignment_with_stellar] Signature date fields already set, no update needed")
            
            return True
        else:
            _logger.warning("[auto_sign_assignment_with_stellar] No files were successfully processed")
            return False

    def extract_text_from_excel_attachment(self, attachment):
        """
        Извлекает текст из Excel файлов (.xlsx, .xls) для анализа YandexGPT
        """
        try:
            if not attachment or not attachment.datas:
                _logger.warning("[extract_text_from_excel_attachment] пустое вложение или отсутствуют данные")
                return None

            # Проверяем доступность openpyxl
            if not OPENPYXL_AVAILABLE:
                _logger.warning("[extract_text_from_excel_attachment] openpyxl не установлен")
                return None

            # Быстрая проверка расширения/типов
            name_lower = (attachment.name or '').lower()
            mimetype = (attachment.mimetype or '').lower()
            
            excel_extensions = ('.xlsx', '.xls')
            excel_mimetypes = (
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel'
            )
            
            if not (name_lower.endswith(excel_extensions) or any(mt in mimetype for mt in excel_mimetypes)):
                _logger.info(f"[extract_text_from_excel_attachment] вложение {attachment.name} не Excel (mimetype={mimetype})")
                return None

            excel_bytes = base64.b64decode(attachment.datas)
            return self._extract_text_from_excel_bytes(excel_bytes, attachment.name)

        except Exception as e:
            _logger.error(f"[extract_text_from_excel_attachment] ошибка обработки {attachment and attachment.name}: {str(e)}")
            return None

    def extract_text_from_doc_attachment(self, attachment):
        """
        Извлекает текст из старых .doc файлов (Microsoft Word 97-2003) для анализа YandexGPT
        Использует LibreOffice для конвертации в текст
        """
        try:
            if not attachment or not attachment.datas:
                _logger.warning("[extract_text_from_doc_attachment] пустое вложение или отсутствуют данные")
                return None

            # Быстрая проверка расширения/типов
            name_lower = (attachment.name or '').lower()
            mimetype = (attachment.mimetype or '').lower()
            
            # Проверяем что это именно .doc файл (старый формат)
            if not (name_lower.endswith('.doc') or mimetype == 'application/msword'):
                _logger.info(f"[extract_text_from_doc_attachment] вложение {attachment.name} не DOC (mimetype={mimetype})")
                return None

            # Исключаем .docx файлы (они обрабатываются отдельно)
            if name_lower.endswith('.docx') or 'officedocument' in mimetype:
                _logger.info(f"[extract_text_from_doc_attachment] вложение {attachment.name} является DOCX, а не DOC")
                return None

            doc_bytes = base64.b64decode(attachment.datas)
            return self._extract_text_from_doc_using_libreoffice(doc_bytes, attachment.name)

        except Exception as e:
            _logger.error(f"[extract_text_from_doc_attachment] ошибка обработки {attachment and attachment.name}: {str(e)}")
            return None

    def _extract_text_from_doc_using_libreoffice(self, doc_bytes, filename=None):
        """
        Извлекает текст из .doc файла используя LibreOffice для конвертации в текст
        """
        if not doc_bytes:
            return None

        try:
            # Создаем временный .doc файл
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as doc_temp:
                doc_temp.write(doc_bytes)
                doc_path = doc_temp.name

            # Создаем временную директорию для выходного файла
            out_dir = tempfile.mkdtemp()
            
            try:
                # Конвертируем .doc в .txt используя LibreOffice
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', out_dir, doc_path
                ]
                _logger.info(f"[_extract_text_from_doc_using_libreoffice] Выполняем команду: {' '.join(cmd)}")
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                if result.returncode != 0:
                    _logger.error(f"[_extract_text_from_doc_using_libreoffice] Ошибка LibreOffice: {result.stderr}")
                    return None

                # Ищем созданный .txt файл
                txt_filename = os.path.splitext(os.path.basename(doc_path))[0] + '.txt'
                txt_path = os.path.join(out_dir, txt_filename)
                
                if not os.path.exists(txt_path):
                    _logger.error(f"[_extract_text_from_doc_using_libreoffice] Конвертированный TXT файл не найден: {txt_path}")
                    return None

                # Читаем текст из файла
                with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                
                if not text or not text.strip():
                    _logger.warning(f"[_extract_text_from_doc_using_libreoffice] Извлеченный текст пуст для файла {filename}")
                    return None
                
                # Очищаем текст от лишних пробелов и переносов
                text = re.sub(r'\n{3,}', '\n\n', text)  # Убираем множественные переносы строк
                text = re.sub(r'[ \t]+', ' ', text)     # Нормализуем пробелы
                text = text.strip()
                
                _logger.info(f"[_extract_text_from_doc_using_libreoffice] Успешно извлечен текст из {filename}: {len(text)} символов")
                return text if text else None
                
            finally:
                # Очищаем временные файлы
                try:
                    os.unlink(doc_path)
                except Exception:
                    pass
                try:
                    txt_filename = os.path.splitext(os.path.basename(doc_path))[0] + '.txt'
                    txt_path_cleanup = os.path.join(out_dir, txt_filename)
                    if os.path.exists(txt_path_cleanup):
                        os.unlink(txt_path_cleanup)
                    os.rmdir(out_dir)
                except Exception:
                    pass
                    
        except Exception as e:
            _logger.error(f"[_extract_text_from_doc_using_libreoffice] Ошибка при извлечении текста из DOC файла: {str(e)}")
            return None

    @staticmethod
    def _extract_text_from_excel_bytes(excel_bytes, filename=None):
        """
        Извлекает текст из Excel файла в виде байт
        """
        if not excel_bytes or not OPENPYXL_AVAILABLE:
            return None

        try:
            # Создаем временный файл для openpyxl
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
                temp_file.write(excel_bytes)
                temp_path = temp_file.name

            try:
                # Открываем файл в read-only режиме для производительности
                workbook = load_workbook(temp_path, read_only=True, data_only=True)
                
                all_text_parts = []
                
                # Обрабатываем все листы
                for sheet_name in workbook.sheetnames:
                    try:
                        worksheet = workbook[sheet_name]
                        sheet_parts = [f"=== Лист: {sheet_name} ==="]
                        
                        # Собираем данные построчно
                        for row_index, row in enumerate(worksheet.iter_rows(values_only=True), 1):
                            if not row or all(cell is None or str(cell).strip() == '' for cell in row):
                                continue  # Пропускаем пустые строки
                            
                            row_parts = []
                            for col_index, cell_value in enumerate(row, 1):
                                if cell_value is not None and str(cell_value).strip():
                                    # Форматируем значение для читаемости
                                    formatted_value = str(cell_value).strip()
                                    row_parts.append(f"[{col_index}]: {formatted_value}")
                            
                            if row_parts:
                                sheet_parts.append(f"Строка {row_index}: " + " | ".join(row_parts))
                        
                        if len(sheet_parts) > 1:  # Есть данные кроме заголовка листа
                            all_text_parts.extend(sheet_parts)
                            all_text_parts.append("")  # Пустая строка между листами
                            
                    except Exception as e:
                        _logger.warning(f"[_extract_text_from_excel_bytes] ошибка обработки листа '{sheet_name}': {str(e)}")
                        continue
                
                workbook.close()
                
                if not all_text_parts:
                    return None
                
                # Склеиваем все части
                text = '\n'.join(all_text_parts)
                
                # Чистим лишние пустые строки
                text = re.sub(r'\n{3,}', '\n\n', text)
                return text.strip() if text.strip() else None
                
            finally:
                # Удаляем временный файл
                try:
                    os.unlink(temp_path)
                except Exception:
                    pass
                    
        except Exception as e:
            _logger.error(f"[_extract_text_from_excel_bytes] ошибка извлечения текста из Excel: {str(e)}")
            return None

    def _find_agent_signature_position_in_pdf(self, pdf_bytes):
        """
        Ищет места для подписи агента в PDF документе "Индивидуал".
        Ищет ТОЛЬКО конкретные блоки с "АГЕНТ" + "Подпись:" или "THE AGENT" + "By:".
        Возвращает список позиций для подписи.
        """
        if not PYMUPDF_AVAILABLE or not pdf_bytes:
            return []
        
        matches = []
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            
            for page_index in range(len(doc)):
                page = doc[page_index]
                text_dict = page.get_text('dict')
                
                # Собираем все блоки текста на странице
                page_blocks = []
                for block in text_dict.get('blocks', []):
                    block_text = ''
                    block_bbox = None
                    
                    for line in block.get('lines', []):
                        for span in line.get('spans', []):
                            span_text = (span.get('text') or '').strip()
                            if span_text:
                                block_text += span_text + ' '
                                if block_bbox is None:
                                    block_bbox = span.get('bbox')
                    
                    if block_text.strip():
                        page_blocks.append({
                            'text': block_text.strip(),
                            'bbox': block_bbox
                        })
                
                # Собираем весь текст страницы для поиска контекста
                page_text = ' '.join([block['text'] for block in page_blocks])
                
                # Ищем ТОЛЬКО конкретные паттерны подписи
                for block in page_blocks:
                    block_text = block['text']
                    
                    # Логируем все блоки для отладки
                    if any(word in block_text.upper() for word in ['АГЕНТ', 'THE AGENT', 'ПОДПИСЬ', 'BY:']):
                        _logger.info(f"[DEBUG] Found potential signature block: '{block_text[:150]}'")
                    
                    # СТРОГИЕ паттерны - ищем только блоки подписи в конце документа
                    
                    # Паттерн 1: Блок "ПРИНЦИПАЛ" с подписью (пропускаем)
                    if 'ПРИНЦИПАЛ' in block_text and ('Подпись:' in block_text or 'By:' in block_text):
                        _logger.info(f"[SKIP] Skipping PRINCIPAL signature block: '{block_text[:50]}...'")
                        continue
                    
                    # Паттерн 2: Русский блок агента - ТОЛЬКО если есть "АГЕНТ" + "Подпись:" + подчеркивания + "МП"
                    russian_agent_signature = (
                        'АГЕНТ' in block_text and 
                        'Подпись:' in block_text and 
                        '_' in block_text and
                        'МП' in block_text and
                        'Директор' in block_text  # Дополнительная проверка
                    )
                    
                    # Паттерн 3: Английский блок агента - ТОЛЬКО если есть "THE AGENT" + "By:" + подчеркивания + "Director"
                    english_agent_signature = (
                        'THE AGENT' in block_text and 
                        'By:' in block_text and 
                        '_' in block_text and
                        'Director' in block_text  # Дополнительная проверка
                    )
                    
                    if russian_agent_signature or english_agent_signature:
                        _logger.info(f"[_find_agent_signature_position_in_pdf] Found AGENT signature block: '{block_text[:100]}...'")
                        _logger.info(f"[_find_agent_signature_position_in_pdf] Patterns - Russian: {russian_agent_signature}, English: {english_agent_signature}")
                        matches.append({
                            'page_number': page_index + 1,
                            'bbox': tuple(block['bbox']),
                            'text': block_text,
                            'is_russian': russian_agent_signature,
                            'is_english': english_agent_signature
                        })
            
            doc.close()
            _logger.info(f"[_find_agent_signature_position_in_pdf] Found {len(matches)} agent signature positions")
            
            # Если ничего не найдено, попробуем старую логику как fallback
            if not matches:
                _logger.warning("[_find_agent_signature_position_in_pdf] No matches with new logic, trying fallback...")
                matches = self._find_agent_signature_position_fallback(pdf_bytes)
            
        except Exception as e:
            _logger.error(f"[_find_agent_signature_position_in_pdf] Error searching for agent signature position: {e}")
        
        return matches
    
    def _find_agent_signature_position_fallback(self, pdf_bytes):
        """Fallback метод поиска позиций для подписи - старая логика"""
        if not PYMUPDF_AVAILABLE or not pdf_bytes:
            return []
        
        matches = []
        try:
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            
            for page_index in range(len(doc)):
                page = doc[page_index]
                text_dict = page.get_text('dict')
                
                for block in text_dict.get('blocks', []):
                    for line in block.get('lines', []):
                        # Собираем весь текст строки
                        line_text = ''
                        line_bbox = None
                        
                        for span in line.get('spans', []):
                            span_text = (span.get('text') or '').strip()
                            if span_text:
                                line_text += span_text + ' '
                                if line_bbox is None:
                                    line_bbox = span.get('bbox')
                        
                        line_text = line_text.strip()
                        
                        # БОЛЕЕ ИЗБИРАТЕЛЬНАЯ FALLBACK ЛОГИКА
                        # Ищем только строки с подписями агента, исключая принципала
                        
                        # Пропускаем строки с ПРИНЦИПАЛ
                        if 'ПРИНЦИПАЛ' in line_text or 'PRINCIPAL' in line_text:
                            continue
                            
                        # Пропускаем заголовки и общие упоминания
                        if 'АГЕНТСКОМУ ДОГОВОРУ' in line_text or 'AGENCY CONTRACT' in line_text:
                            continue
                        
                        # Ищем конкретные паттерны подписи агента
                        russian_agent_line = (
                            'АГЕНТ' in line_text and 
                            ('Подпись:' in line_text or ('_' in line_text and 'МП' in line_text))
                        )
                        
                        english_agent_line = (
                            'THE AGENT' in line_text and 
                            ('By:' in line_text or '_' in line_text)
                        )
                        
                        # Отдельные строки с подписями (только если рядом нет ПРИНЦИПАЛ)
                        signature_line = (
                            ('Подпись:' in line_text or 'By:' in line_text) and 
                            '_' in line_text and
                            'ПРИНЦИПАЛ' not in line_text and 'PRINCIPAL' not in line_text
                        )
                        
                        # Ищем строки только с "АГЕНТ" (для русского) или "THE AGENT" (для английского)
                        standalone_agent_russian = (
                            line_text.strip() == 'АГЕНТ' or 
                            (line_text.strip().startswith('АГЕНТ') and len(line_text.strip()) < 10)
                        )
                        
                        standalone_agent_english = (
                            line_text.strip() == 'THE AGENT' or
                            (line_text.strip().startswith('THE AGENT') and len(line_text.strip()) < 15)
                        )
                        
                        if russian_agent_line or english_agent_line or signature_line or standalone_agent_russian or standalone_agent_english:
                            _logger.info(f"[FALLBACK] Found agent signature line: '{line_text}'")
                            if line_bbox:
                                matches.append({
                                    'page_number': page_index + 1,
                                    'bbox': tuple(line_bbox),
                                    'text': line_text,
                                    'is_russian': russian_agent_line or (signature_line and 'Подпись:' in line_text) or standalone_agent_russian,
                                    'is_english': english_agent_line or (signature_line and 'By:' in line_text) or standalone_agent_english
                                })
            
            doc.close()
            _logger.info(f"[FALLBACK] Found {len(matches)} signature positions")
            
        except Exception as e:
            _logger.error(f"[FALLBACK] Error: {e}")
        
        return matches

    def _detect_agent_type_from_record(self):
        """
        Определяет тип агента из записи заявки для документа "Индивидуал".
        Возвращает 'ТДК', 'ИНДО ТРЕЙД' или 'СТЕЛЛАР'.
        """
        try:
            # Получаем название агента из записи
            agent_name = ''
            if self.agent_id and self.agent_id.name:
                agent_name = self.agent_id.name.upper()
            
            _logger.info(f"[_detect_agent_type_from_record] Agent name from record: '{agent_name}'")
            
            # Определяем тип агента по названию
            if any(keyword in agent_name for keyword in ['ТДК', 'TDK']):
                return 'ТДК'
            elif any(keyword in agent_name for keyword in ['ИНДО', 'INDO', 'ТРЕЙД', 'TRADE']):
                return 'ИНДО ТРЕЙД'
            elif any(keyword in agent_name for keyword in ['СТЕЛЛАР', 'STELLAR']):
                return 'СТЕЛЛАР'
            else:
                _logger.info(f"[_detect_agent_type_from_record] Unknown agent type, defaulting to СТЕЛЛАР")
                return 'СТЕЛЛАР'
                
        except Exception as e:
            _logger.error(f"[_detect_agent_type_from_record] Error detecting agent type: {e}")
            return 'СТЕЛЛАР'

    def _find_agent_signature_and_stamp_records(self, agent_type):
        """
        Находит записи подписи и печати для указанного типа агента.
        Возвращает кортеж (signature_record, stamp_record).
        """
        if agent_type == 'ТДК':
            search_terms = ['ТДК', 'TDK']
        elif agent_type == 'ИНДО ТРЕЙД':
            search_terms = ['ИНОТРЕЙД', 'INDOTRADE', 'INDO']
        else:  # СТЕЛЛАР
            search_terms = ['СТЕЛЛАР', 'STELLAR']
        
        sig_record = None
        stamp_record = None
        
        # Ищем подпись
        for term in search_terms:
            sig_record = self.env['signature.library'].search([
                ('name', 'ilike', term),
                ('signature_type', '=', 'signature'),
                ('active', '=', True)
            ], limit=1)
            if sig_record:
                break
        
        # Ищем печать
        for term in search_terms:
            stamp_record = self.env['signature.library'].search([
                ('name', 'ilike', term),
                ('signature_type', '=', 'stamp'),
                ('active', '=', True)
            ], limit=1)
            if stamp_record:
                break
        
        _logger.info(f"[_find_agent_signature_and_stamp_records] Found signature: {sig_record.name if sig_record else 'None'}")
        _logger.info(f"[_find_agent_signature_and_stamp_records] Found stamp: {stamp_record.name if stamp_record else 'None'}")
        
        return sig_record, stamp_record

    def _sign_individual_document(self, pdf_bytes, agent_type):
        """
        Подписывает PDF документ "Индивидуал" подписью и печатью агента.
        Ищет места с "АГЕНТ"/"THE AGENT" и "Подпись:"/"By:" и ставит подписи правее.
        """
        _logger.info(f"[_sign_individual_document] Starting signature process for agent type: {agent_type}")
        
        if not PYMUPDF_AVAILABLE or not pdf_bytes:
            _logger.warning("[_sign_individual_document] PyMuPDF not available or no PDF data")
            return pdf_bytes
        
        try:
            # Находим позиции для подписи
            _logger.info("[_sign_individual_document] Searching for signature positions...")
            matches = self._find_agent_signature_position_in_pdf(pdf_bytes)
            _logger.info(f"[_sign_individual_document] Found {len(matches)} signature positions")
            
            if not matches:
                _logger.warning("[_sign_individual_document] No agent signature positions found - returning unsigned PDF")
                return pdf_bytes
            
            # Получаем записи подписи и печати
            sig_record, stamp_record = self._find_agent_signature_and_stamp_records(agent_type)
            if not sig_record or not stamp_record:
                _logger.warning(f"[_sign_individual_document] {agent_type} signature/stamp not found")
                return pdf_bytes
            
            doc = pymupdf.open(stream=pdf_bytes, filetype='pdf')
            
            # Группируем совпадения по страницам
            pages_to_sign = {}
            for match in matches:
                page_num = match['page_number'] - 1
                if page_num not in pages_to_sign:
                    pages_to_sign[page_num] = []
                pages_to_sign[page_num].append(match)
            
            # Подписываем каждую страницу
            for page_num, page_matches in pages_to_sign.items():
                page = doc[page_num]
                
                # Обрабатываем каждый найденный блок подписи
                for match in page_matches:
                    text = match['text']
                    bbox = match['bbox']
                    
                    _logger.info(f"[_sign_individual_document] Processing signature block: '{text[:50]}...'")
                    
                    # Размеры подписи и печати
                    sig_w = sig_record.default_width or 120
                    sig_h = sig_record.default_height or 40
                    stamp_w = stamp_record.default_width or 80
                    stamp_h = stamp_record.default_height or 80
                    
                    # Определяем позицию для подписи и печати более точно
                    if match.get('is_russian'):
                        if 'Подпись:' in text:
                            # Строка с "Подпись: _______"
                            sig_x = bbox[0] + 80   # После "Подпись:"
                            sig_y = bbox[1] - 5    # На той же высоте
                            
                            # Печать правее подписи, где должно быть "МП"
                            stamp_x = bbox[0] + 250  # Где обычно "МП"
                            stamp_y = bbox[1] - 10   # Немного выше для центрирования
                        else:
                            # Строка только с "АГЕНТ" - ищем место для подписи ниже
                            sig_x = bbox[0] + 80   # Отступ от левого края
                            sig_y = bbox[3] + 10   # Ниже строки "АГЕНТ"
                            
                            # Печать правее
                            stamp_x = sig_x + sig_w + 30
                            stamp_y = sig_y - 5
                        
                    elif match.get('is_english'):
                        if 'By:' in text:
                            # Строка с "By: _______"
                            sig_x = bbox[0] + 50   # После "By:"
                            sig_y = bbox[1] - 5    # На той же высоте
                            
                            # Печать правее подписи, где должно быть "Stamp"
                            stamp_x = bbox[0] + 200  # Где обычно "Stamp"
                            stamp_y = bbox[1] - 10   # Немного выше для центрирования
                        else:
                            # Строка только с "THE AGENT" - ищем место для подписи ниже
                            sig_x = bbox[0] + 50   # Отступ от левого края
                            sig_y = bbox[3] + 10   # Ниже строки "THE AGENT"
                            
                            # Печать правее
                            stamp_x = sig_x + sig_w + 30
                            stamp_y = sig_y - 5
                    
                    else:
                        # Fallback позиция
                        sig_x = bbox[0] + 50
                        sig_y = bbox[1]
                        stamp_x = sig_x + sig_w + 20
                        stamp_y = sig_y - 5
                    
                    # Логируем детали позиционирования
                    _logger.info(f"[_sign_individual_document] Text: '{text[:50]}...'")
                    _logger.info(f"[_sign_individual_document] BBox: {bbox}")
                    _logger.info(f"[_sign_individual_document] Signature position: ({sig_x}, {sig_y}) size: {sig_w}x{sig_h}")
                    _logger.info(f"[_sign_individual_document] Stamp position: ({stamp_x}, {stamp_y}) size: {stamp_w}x{stamp_h}")
                    
                    # Проставляем подпись
                    sig_rect = pymupdf.Rect(sig_x, sig_y, sig_x + sig_w, sig_y + sig_h)
                    page.insert_image(sig_rect, stream=base64.b64decode(sig_record.image))
                    
                    # Проставляем печать
                    stamp_rect = pymupdf.Rect(stamp_x, stamp_y, stamp_x + stamp_w, stamp_y + stamp_h)
                    page.insert_image(stamp_rect, stream=base64.b64decode(stamp_record.image))
                    
                    _logger.info(f"[_sign_individual_document] Added {agent_type} signature and stamp on page {page_num + 1}")
            
            # Сохраняем подписанный PDF
            signed_pdf_bytes = doc.write()
            doc.close()
            
            _logger.info(f"[_sign_individual_document] Successfully signed individual document with {agent_type} signatures")
            return signed_pdf_bytes
            
        except Exception as e:
            _logger.error(f"[_sign_individual_document] Error signing individual document: {e}")
            return pdf_bytes